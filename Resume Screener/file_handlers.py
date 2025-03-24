"""
File handling module for Resume Screener.

Provides functions for reading and extracting text from various document formats.
"""

import os
import logging
import glob
from pathlib import Path
from typing import List, Optional

import docx
import PyPDF2

logger = logging.getLogger(__name__)


def validate_job_description_path(file_path: str) -> bool:
    """
    Validate if the given file path is a valid job description file.

    Args:
        file_path: Path to the job description file

    Returns:
        bool: True if valid, False otherwise
    """
    if not os.path.isfile(file_path):
        logger.warning(f"File does not exist: {file_path}")
        return False

    # Check file extension
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in [".txt", ".docx"]:
        logger.warning(f"Unsupported file format: {ext} for job description")
        return False

    return True


def extract_job_description_text(file_path: str) -> Optional[str]:
    """
    Extract text from job description file.

    Args:
        file_path: Path to the job description file

    Returns:
        str: Extracted text or None if extraction failed
    """
    try:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".txt":
            # Read text file
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".docx":
            # Read DOCX file using python-docx
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])

        else:
            logger.error(f"Unsupported file format: {ext}")
            return None

    except Exception as e:
        logger.exception(f"Error extracting text from {file_path}: {str(e)}")
        return None


def get_resume_files(directory: str) -> List[str]:
    """
    Get a list of resume files in the specified directory.

    Args:
        directory: Directory to search for resume files

    Returns:
        List[str]: List of absolute paths to resume files
    """
    resume_files = []

    # Get DOCX files
    docx_files = glob.glob(os.path.join(directory, "**", "*.docx"), recursive=True)
    resume_files.extend(docx_files)

    # Get PDF files
    pdf_files = glob.glob(os.path.join(directory, "**", "*.pdf"), recursive=True)
    resume_files.extend(pdf_files)

    logger.info(
        f"Found {len(resume_files)} resume files: {len(docx_files)} DOCX, {len(pdf_files)} PDF"
    )

    return resume_files


def extract_resume_text(file_path: str) -> Optional[str]:
    """
    Extract text from a resume file.

    Args:
        file_path: Path to the resume file

    Returns:
        str: Extracted text or None if extraction failed
    """
    try:
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            # Read DOCX file using python-docx
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.debug(f"Successfully extracted text from DOCX file: {file_path}")
            return text

        elif ext == ".pdf":
            # Read PDF file using PyPDF2
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text() or ""

                logger.debug(f"Successfully extracted text from PDF file: {file_path}")
                return text

        else:
            logger.error(f"Unsupported file format: {ext}")
            return None

    except Exception as e:
        logger.exception(f"Error extracting text from {file_path}: {str(e)}")
        return None
